from docx import Document
from docx.shared import Inches
from pathlib import Path
import re

root = Path('.')
md_path = root / '论文.md'
doc_path = root / '论文.docx'

image_paths = {
    'output/acc_compare.png': root / 'output' / 'acc_compare.png',
    'output/noise_robust.png': root / 'output' / 'noise_robust.png',
    'output/overfit_gap.png': root / 'output' / 'overfit_gap.png',
    'output/eda/label_dist.png': root / 'output' / 'eda' / 'label_dist.png',
    'output/eda/corr_heat.png': root / 'output' / 'eda' / 'corr_heat.png',
    'output/eda/outlier_box.png': root / 'output' / 'eda' / 'outlier_box.png',
    'output/loss_curves.png': root / 'output' / 'loss_curves.png'
}

text = md_path.read_text(encoding='utf-8')
lines = text.splitlines()

doc = Document()

in_table = False
table_header = None
table_rows = []

def flush_table():
    global in_table, table_header, table_rows
    if not in_table:
        return
    if table_header is None or not table_rows:
        in_table = False
        table_header = None
        table_rows = []
        return
    table = doc.add_table(rows=1, cols=len(table_header))
    hdr_cells = table.rows[0].cells
    for i, heading in enumerate(table_header):
        hdr_cells[i].text = heading.strip()
    for row in table_rows:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = cell.strip()
    doc.add_paragraph('')
    in_table = False
    table_header = None
    table_rows = []

for line in lines:
    if in_table:
        if line.startswith('|') and '---' not in line:
            cols = [c.strip() for c in line.strip().strip('|').split('|')]
            table_rows.append(cols)
            continue
        else:
            flush_table()
    if line.startswith('# '):
        doc.add_heading(line[2:].strip(), level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
    elif re.match(r'!\[[^\]]*\]\(([^)]+)\)', line):
        match = re.match(r'!\[[^\]]*\]\(([^)]+)\)', line)
        image_key = match.group(1)
        if image_key in image_paths and image_paths[image_key].exists():
            doc.add_picture(str(image_paths[image_key]), width=Inches(6))
        else:
            doc.add_paragraph(line)
    elif line.startswith('|'):
        if '---' in line:
            continue
        if table_header is None:
            table_header = [c.strip() for c in line.strip().strip('|').split('|')]
            in_table = True
            table_rows = []
            continue
        else:
            cols = [c.strip() for c in line.strip().strip('|').split('|')]
            table_rows.append(cols)
    else:
        doc.add_paragraph(line)

flush_table()
doc.save(doc_path)
print('Saved', doc_path)
