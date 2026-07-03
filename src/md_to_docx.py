from docx import Document
from docx.shared import Inches
from pathlib import Path
import csv
import re

md_path = Path('..') / '论文.md'
doc_path = Path('..') / '论文.docx'

image_paths = {
    'output/acc_compare.png': Path('..') / 'output' / 'acc_compare.png',
    'output/noise_robust.png': Path('..') / 'output' / 'noise_robust.png',
    'output/overfit_gap.png': Path('..') / 'output' / 'overfit_gap.png',
    'output/eda/label_dist.png': Path('..') / 'output' / 'eda' / 'label_dist.png',
    'output/eda/corr_heat.png': Path('..') / 'output' / 'eda' / 'corr_heat.png',
    'output/eda/outlier_box.png': Path('..') / 'output' / 'eda' / 'outlier_box.png'
}

text = md_path.read_text(encoding='utf-8')

doc = Document()
for line in text.splitlines():
    if line.startswith('# '):
        doc.add_heading(line[2:].strip(), level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
    elif line.startswith('!['):
        match = re.match(r'!\[[^\]]*\]\(([^)]+)\)', line)
        if match:
            image_key = match.group(1)
            if image_key in image_paths:
                img_path = image_paths[image_key]
                if img_path.exists():
                    doc.add_picture(str(img_path), width=Inches(6))
                    continue
        doc.add_paragraph(line)
    elif line.startswith('- `') and line.endswith('`'):
        image_key = line[3:-1]
        if image_key in image_paths:
            img_path = image_paths[image_key]
            if img_path.exists():
                doc.add_paragraph(image_key)
                doc.add_picture(str(img_path), width=Inches(6))
                continue
        doc.add_paragraph(line)
    elif line.startswith('|') and '---' not in line:
        # skip markdown table rows; use CSV table instead
        continue
    else:
        doc.add_paragraph(line)

# 追加模型对比表格
csv_path = Path('..') / 'output' / 'model_compare.csv'
if csv_path.exists():
    with csv_path.open(newline='', encoding='utf-8-sig') as f:
        reader = csv.reader(f)
        rows = list(reader)
    if rows:
        doc.add_page_break()
        doc.add_heading('模型对比表', level=2)
        table = doc.add_table(rows=1, cols=len(rows[0]))
        hdr_cells = table.rows[0].cells
        for i, heading in enumerate(rows[0]):
            hdr_cells[i].text = heading
        for row in rows[1:]:
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = value

doc.save(doc_path)
print('Saved', doc_path)
